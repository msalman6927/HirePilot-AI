import requests
import docx
import json
import time
import os

def create_sample_docx(filename="test_cv.docx"):
    doc = docx.Document()
    doc.add_heading('Jane Doe', 0)
    doc.add_paragraph('jane.doe@example.com | 123-456-7890')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Skilled Software Developer with expertise in Python and AI.')
    
    doc.add_heading('Experience', level=1)
    p = doc.add_paragraph()
    p.add_run('Senior Developer | TechCorp | 2021-Present').bold = True
    doc.add_paragraph('Leading backend development teams.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, Docker, Kubernetes, AWS, Machine Learning')
    
    doc.save(filename)
    return filename

def test_upload():
    # Helper to check health
    url_base = "http://127.0.0.1:8000"
    try:
        health = requests.get(f"{url_base}/")
        print(f"Health Check (/): {health.status_code} {health.text}")
    except Exception as e:
        print(f"Health Check failed: {e}")
        return

    # Helper to check CV router existence via docs
    try:
        docs = requests.get(f"{url_base}/openapi.json")
        if docs.status_code == 200:
            paths = docs.json().get("paths", {}).keys()
            print("Available Paths:")
            for p in paths:
                print(f" - {p}")
            if "/cv/upload" not in paths:
                print("ERROR: /cv/upload not found in OpenAPI spec!")
        else:
            print(f"Docs check failed: {docs.status_code}")
    except Exception as e:
        print(f"Docs check failed: {e}")

    filename = create_sample_docx()
    url = f"{url_base}/cv/upload"
    
    print(f"\nUploading {filename} to {url}...")
    
    try:
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(url, files=files)
            
        print(f"Status Code: {response.status_code}")
        if response.status_code == 200:
            print("Response JSON:")
            print(json.dumps(response.json(), indent=2))
        else:
            print("Response Text:")
            print(response.text)
            
    except Exception as e:
        print(f"Error during upload: {e}")

if __name__ == "__main__":
    time.sleep(5) 
    test_upload()
